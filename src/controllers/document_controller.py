import os
import tempfile
import traceback
import uuid

import openai
import PyPDF2
import requests
from docx import Document
from flask import jsonify, request
from PyPDF2 import PdfReader

from db.firebase import bucket
from src.config.openai_config import openai
from src.config.pinecone_config import ServerlessSpec, pc, pinecone_index

embedding_model = 'text-embedding-ada-002'

def extract_text_metadata(file_url):
    file_ext = os.path.splitext(file_url)[1].lower()
    
    # Download the file content from the URL
    response = requests.get(file_url)
    if response.status_code != 200:
        raise Exception(f"Failed to download file from {file_url}")
    
    # Handle DOCX files
    if file_ext == '.docx':
        from io import BytesIO
        file_stream = BytesIO(response.content)
        document = Document(file_stream)
        content = '\n'.join([para.text for para in document.paragraphs])
        metadata = {
            "file_name": os.path.basename(file_url),
            "file_path": file_url,
            "file_type": "DOCX"
        }
        return content, metadata
    
    # Handle PDF files
    elif file_ext == '.pdf':
        from io import BytesIO
        file_stream = BytesIO(response.content)
        reader = PdfReader(file_stream)
        content = ''
        for page in reader.pages:
            text = page.extract_text()
            if text:
                content += text
        metadata = {
            "file_name": os.path.basename(file_url),
            "file_path": file_url,
            "file_type": "PDF",
            "num_pages": len(reader.pages)
        }
        return content, metadata
    
    # Handle TXT files
    elif file_ext == '.txt':
        content = response.text  # No need for further processing for plain text files
        metadata = {
            "file_name": os.path.basename(file_url),
            "file_path": file_url,
            "file_type": "TXT",
            "num_characters": len(content)
        }
        return content, metadata
    
    else:
        raise ValueError("Unsupported file format: Only DOCX, PDF, and TXT are supported.")

def upload_and_embed_document(content, metadata):
    """
    Create an embedding for the content and upload it to Pinecone.
    """
    response = openai.Embedding.create(
        input=content,
        model=embedding_model
    )
    embedding_vector = response['data'][0]['embedding']

    document_id = str(uuid.uuid4())
    pinecone_index.upsert(vectors=[(document_id, embedding_vector, metadata)])
    return document_id

def upload_document():
    """
    Flask route to handle document uploads.
    """
    if 'documents' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    files = request.files.getlist('documents')
    if not files or files[0].filename == '':
        return jsonify({"error": "No files selected for uploading"}), 400

    document_ids = []

    for file in files:
        try:
            # Generate a unique identifier for the file
            unique_id = str(uuid.uuid4())
            filename = file.filename
            file_ext = os.path.splitext(filename)[1].lower()
            blob_name = f"uploaded_documents/{unique_id}_{filename}"
            blob = bucket.blob(blob_name)

            # Upload the file to Firebase Storage
            blob.upload_from_file(file.stream, content_type=file.content_type)
            # Optionally, make the file public
            blob.make_public()

            # Get the public URL of the uploaded file
            file_url = blob.public_url

            print(f"File uploaded successfully to {file_url}")  # Debugging line

            # **Download the file from the URL to a temporary local path**
            response = requests.get(file_url)
            if response.status_code != 200:
                raise Exception(f"Failed to download file from {file_url}")

            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(response.content)
                temp_file_path = temp_file.name

            content, metadata = extract_text_metadata(file_url)

            # Upload and embed the document
            document_id = upload_and_embed_document(content, metadata)
            document_ids.append({
                "document_id": document_id,
                "file_name": metadata["file_name"],
                "file_path": file_url
            })

        except Exception as e:
            print(f"Failed to process file {file.filename}: {e}")  # Debugging line
            traceback.print_exc()
            return jsonify({"error": f"Failed to process file {file.filename}: {e}"}), 500

    return jsonify({"status": "success", "uploaded_documents": document_ids}), 200

# import os
# import tempfile
# import traceback
# import uuid

# import openai
# import PyPDF2
# from docx import Document
# from flask import jsonify, request
# from PyPDF2 import PdfReader

# from src.config.openai_config import openai
# from src.config.pinecone_config import ServerlessSpec, pc, pinecone_index

# embedding_model = 'text-embedding-ada-002'

# def extract_text_metadata(file_path):
#     file_ext = os.path.splitext(file_path)[1].lower()
    
#     # Handle DOCX files
#     if file_ext == '.docx':
#         document = Document(file_path)
#         content = '\n'.join([para.text for para in document.paragraphs])
#         metadata = {
#             "file_name": os.path.basename(file_path),
#             "file_path": file_path,
#             "file_type": "DOCX"
#         }
#         return content, metadata
    
#     # Handle PDF files
#     elif file_ext == '.pdf':
#         reader = PdfReader(file_path)
#         content = ''
#         for page in reader.pages:
#             content += page.extract_text()  # Extract text from each page
        
#         metadata = {
#             "file_name": os.path.basename(file_path),
#             "file_path": file_path,
#             "file_type": "PDF",
#             "num_pages": len(reader.pages)
#         }
#         return content, metadata
    
#     # Handle TXT files
#     elif file_ext == '.txt':
#         with open(file_path, 'r', encoding='utf-8') as file:
#             content = file.read()
        
#         metadata = {
#             "file_name": os.path.basename(file_path),
#             "file_path": file_path,
#             "file_type": "TXT",
#             "num_characters": len(content)
#         }
#         return content, metadata
    
#     else:
#         raise ValueError("Unsupported file format: Only DOCX, PDF, and TXT are supported.")

# def upload_and_embed_document(content, metadata):
#     response = openai.Embedding.create(
#         input=content,
#         model=embedding_model
#     )
#     embedding_vector = response['data'][0]['embedding']
    
#     document_id = str(uuid.uuid4())
#     pinecone_index.upsert(vectors=[(document_id, embedding_vector, metadata)])
#     return document_id

# def upload_document():
#     if 'documents' not in request.files:
#         return jsonify({"error": "No file part in the request"}), 400

#     files = request.files.getlist('documents')
#     if not files or files[0].filename == '':
#         return jsonify({"error": "No files selected for uploading"}), 400

#     # Define your persistent storage path as an absolute path
#     base_dir = os.path.dirname(os.path.abspath(__file__))
#     storage_path = os.path.join(base_dir, '../../uploaded_documents')

#     # Ensure the directory exists
#     os.makedirs(storage_path, exist_ok=True)

#     document_ids = []

#     for file in files:
#         # Construct a unique file path within the storage directory
#         temp_file_path = os.path.join(storage_path, f"{uuid.uuid4()}_{file.filename}")
#         print(f"Saving file to: {temp_file_path}")  # Debugging line

#         try:
#             # Save the file to the specified path
#             file.save(temp_file_path)
#             print(f"File saved successfully to {temp_file_path}")  # Debugging line
#         except Exception as e:
#             print(f"Failed to save file: {e}")  # Debugging line
#             return jsonify({"error": f"Failed to save file: {e}"}), 500

#         # Extract text and metadata and upload to Pinecone
#         content, metadata = extract_text_metadata(temp_file_path)
#         document_id = upload_and_embed_document(content, metadata)
#         document_ids.append({
#             "document_id": document_id,
#             "file_name": metadata["file_name"],
#             "file_path": temp_file_path
#         })

#     return jsonify({"status": "success", "uploaded_documents": document_ids}), 200

def query_document():
    data = request.get_json()
    query_text = data.get("query")
    document_ids = data.get("document_ids")
    
    if not query_text:
        return jsonify({"error": "Query text is required."}), 400
    if not document_ids or not isinstance(document_ids, list):
        return jsonify({"error": "A list of document_ids is required."}), 400

    response = openai.Embedding.create(
        input=query_text,
        model=embedding_model
    )
    query_embedding = response['data'][0]['embedding']

    results = pinecone_index.query(
        vector=query_embedding,
        top_k=10,
        include_metadata=True
    )

    # Debugging: Print the document IDs from Pinecone matches
    print("Document IDs returned by Pinecone:", [match['id'] for match in results['matches']])

    # Debugging: Add print to see how filtering is working
    filtered_matches = [
        match for match in results['matches'] if match['id'].strip() in [doc_id.strip() for doc_id in document_ids]
    ]
    print("Filtered matches:", filtered_matches)

    if not filtered_matches:
        return jsonify({"error": "No matching documents found for the given IDs."}), 404

    # Initialize context and document names
    context = ''
    document_names = []

    for match in filtered_matches:
        file_name = match['metadata']['file_name']
        file_path = match['metadata']['file_path']
        document_names.append(file_name)

        try:
            # Read and append the content of each document
            content = extract_text_metadata(file_path)
            context += f"\nDocument Name: {file_name}\n{content}\n"
        except Exception as e:
            return jsonify({"error": f"Failed to read file: {e}"}), 500

    # Construct the prompt using all matched document contents
    prompt = f"""
You are an AI assistant that provides answers based on specific documents.

Context:
{context}

Question:
{query_text}

Answer:
"""

    chat_response = openai.ChatCompletion.create(
        model='gpt-3.5-turbo',
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500,
        temperature=0.7,
    )

    answer = chat_response['choices'][0]['message']['content'].strip()
    return jsonify({"answer": answer, "document_names": document_names}), 200




